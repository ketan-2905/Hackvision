# from fastapi import FastAPI, UploadFile, File, HTTPException
# from fastapi.middleware.cors import CORSMiddleware
# import uvicorn
# import fitz  # PyMuPDF
# import docx
# import io

# app = FastAPI(title="Resume Text Extraction Service")

# # CORS configuration - allow Next.js dev server
# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["http://localhost:3000", "http://127.0.0.1:3000"],
#     allow_credentials=True,
#     allow_methods=["*"],
#     allow_headers=["*"],
# )

# @app.get("/health")
# async def health_check():
#     """Health check endpoint"""
#     return {"status": "ok", "service": "text-extraction"}

# @app.post("/extract-text")
# async def extract_text(file: UploadFile = File(...)):
#     """
#     Extract text from PDF or DOCX file.
#     Returns: {"text": str, "success": bool, "filename": str}
#     """
#     try:
#         # Read file content
#         content = await file.read()
#         filename = file.filename or "unknown"
        
#         # Extract based on file type
#         if filename.lower().endswith(".pdf"):
#             doc = fitz.open(stream=content, filetype="pdf")
#             text = ""
#             for page in doc:
#                 text += page.get_text() + "\n"
            
#             # Get metadata before closing
#             page_count = len(doc)
#             doc.close()
            
#             return {
#                 "text": text.strip(),
#                 "success": True,
#                 "filename": filename,
#                 "pages": page_count,
#                 "length": len(text.strip())
#             }
            
#         elif filename.lower().endswith(".docx") or filename.lower().endswith(".doc"):
#             doc = docx.Document(io.BytesIO(content))
#             text = "\n".join([para.text for para in doc.paragraphs])
            
#             return {
#                 "text": text.strip(),
#                 "success": True,
#                 "filename": filename,
#                 "paragraphs": len(doc.paragraphs),
#                 "length": len(text.strip())
#             }
            
#         else:
#             raise HTTPException(
#                 status_code=400,
#                 detail=f"Unsupported file type: {filename}. Only PDF and DOCX are supported."
#             )
            
#     except HTTPException:
#         raise
#     except Exception as e:
#         print(f"Extraction error: {str(e)}")
#         raise HTTPException(
#             status_code=500,
#             detail=f"Text extraction failed: {str(e)}"
#         )

# if __name__ == "__main__":
#     print("=" * 60)
#     print("🚀 Resume Text Extraction Service Starting...")
#     print("=" * 60)
#     print("📍 Listening on: http://127.0.0.1:8001")
#     print("🔗 Health check: http://127.0.0.1:8001/health")
#     print("📄 Endpoint: POST http://127.0.0.1:8001/extract-text")
#     print("=" * 60)
    
#     uvicorn.run(
#         "text_extractor:app",
#         host="127.0.0.1",
#         port=8001,
#         reload=True,
#         log_level="info"
#     )

import os
import io
import hashlib
from datetime import datetime

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
import uvicorn

import fitz  # PyMuPDF
import docx

from sqlalchemy import create_engine, Column, String, Integer, DateTime, UniqueConstraint
from sqlalchemy.orm import declarative_base, sessionmaker

from dotenv import load_dotenv
from web3 import Web3

# -------------------------------------------------
# LOAD ENV (THIS IS WHERE YOUR CREDS COME FROM)
# -------------------------------------------------
load_dotenv()

RPC_URL = os.getenv("RPC_URL")
PRIVATE_KEY = os.getenv("PRIVATE_KEY")
CONTRACT_ADDRESS = os.getenv("CONTRACT_ADDRESS")
WALLET_ADDRESS = os.getenv("WALLET_ADDRESS")

if not RPC_URL or not PRIVATE_KEY or not CONTRACT_ADDRESS:
    raise RuntimeError("Missing blockchain env variables")

# -------------------------------------------------
# BLOCKCHAIN SETUP (READINESS CHECK)
# -------------------------------------------------
w3 = Web3(Web3.HTTPProvider(RPC_URL))

if not w3.is_connected():
    raise RuntimeError("Blockchain RPC not reachable")

# -------------------------------------------------
# DATABASE (SQLITE)
# -------------------------------------------------
DATABASE_URL = "sqlite:///./resumes.db"

engine = create_engine(
    DATABASE_URL,
    connect_args={"check_same_thread": False}
)
SessionLocal = sessionmaker(bind=engine, autocommit=False, autoflush=False)
Base = declarative_base()


class ResumeHash(Base):
    __tablename__ = "resume_hashes"

    id = Column(Integer, primary_key=True)
    filename = Column(String, nullable=False)
    file_hash = Column(String, nullable=False, unique=True, index=True)
    created_at = Column(DateTime, default=datetime.utcnow)

    __table_args__ = (
        UniqueConstraint("file_hash", name="uq_resume_hash"),
    )


Base.metadata.create_all(bind=engine)

# -------------------------------------------------
# APP
# -------------------------------------------------
app = FastAPI(title="Resume Hash + Blockchain Guard")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# -------------------------------------------------
# HELPERS
# -------------------------------------------------
def sha256_hex(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


# -------------------------------------------------
# ROUTES
# -------------------------------------------------
@app.get("/health")
async def health():
    return {
        "status": "ok",
        "chain": w3.eth.chain_id,
        "wallet": WALLET_ADDRESS,
    }


@app.post("/extract-text")
async def extract_text(file: UploadFile = File(...)):
    content = await file.read()
    filename = file.filename or "unknown"

    file_hash = sha256_hex(content)

    db = SessionLocal()

    # 🚫 DUPLICATE CHECK
    if db.query(ResumeHash).filter_by(file_hash=file_hash).first():
        raise HTTPException(
            status_code=409,
            detail="Duplicate resume detected. Same file already uploaded."
        )

    # ✅ STORE HASH
    db.add(ResumeHash(
        filename=filename,
        file_hash=file_hash
    ))
    db.commit()

    # -------- TEXT EXTRACTION --------
    if filename.lower().endswith(".pdf"):
        doc = fitz.open(stream=content, filetype="pdf")
        text = "\n".join(p.get_text() for p in doc)
        pages = len(doc)
        doc.close()

        return {
            "success": True,
            "filename": filename,
            "hash": file_hash,
            "pages": pages,
            "text": text.strip(),
        }

    if filename.lower().endswith((".docx", ".doc")):
        doc = docx.Document(io.BytesIO(content))
        text = "\n".join(p.text for p in doc.paragraphs)

        return {
            "success": True,
            "filename": filename,
            "hash": file_hash,
            "paragraphs": len(doc.paragraphs),
            "text": text.strip(),
        }

    raise HTTPException(400, "Only PDF and DOCX supported")


# -------------------------------------------------
# MAIN
# -------------------------------------------------
if __name__ == "__main__":
    print("🚀 Resume Guard Running")
    print("🔐 Blockchain RPC:", RPC_URL)
    print("📍 Wallet:", WALLET_ADDRESS)

    uvicorn.run(
        "text_extractor:app",
        host="127.0.0.1",
        port=8001,
        reload=True
    )